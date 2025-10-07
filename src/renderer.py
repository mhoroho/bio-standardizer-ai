# src/renderer.py
from __future__ import annotations

import json
from pathlib import Path
from typing import Dict, List, Any, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from schema import Bio


# ---------- plan I/O ----------
def load_plan(path: Path) -> Dict[str, Any]:
    return json.loads(Path(path).read_text(encoding="utf-8"))


# ---------- safe helpers ----------
def _safe_set_style(p, style_name: Optional[str]):
    try:
        if style_name:
            p.style = style_name
    except Exception:
        # If the style doesn't exist in this Word template, just skip
        pass


def _add_heading(doc: Document, text: str, style_name: Optional[str]):
    p = doc.add_paragraph(text)
    _safe_set_style(p, style_name or "Heading 2")
    return p


def _add_paragraph(doc: Document, text: str, style_name: Optional[str]):
    p = doc.add_paragraph(text)
    _safe_set_style(p, style_name or "Normal")
    return p


def _add_bullets(
    doc: Document,
    items: List[str],
    style_name: Optional[str],
    limit: Optional[int] = None,
):
    count = 0
    for it in items:
        if limit is not None and count >= limit:
            break
        p = doc.add_paragraph(it)
        # Try the specified bullet style; if it fails, fall back to common names
        if style_name:
            _safe_set_style(p, style_name)
        else:
            # Try common bullet styles
            for candidate in ("List Bullet", "List Paragraph", "Normal"):
                _safe_set_style(p, candidate)
                # If the first one doesn't exist, the next attempt will override
        count += 1


# ---------- header rendering ----------
def _render_header(doc: Document, bio: Bio, header_plan: Dict[str, Any]):
    """
    Render the header block (identity + contacts) using LLM-chosen styles and order.
    header_plan example:
    {
      "name_style": "Heading 1",
      "title_style": "Subtitle",
      "department_style": "Normal",
      "location_style": "Normal",
      "contact_style": "Normal",
      "line_order": ["name","title","department_or_practice","location","contact"]
    }
    """
    name_style = header_plan.get("name_style", "Heading 1")
    title_style = header_plan.get("title_style", "Subtitle")
    dept_style = header_plan.get("department_style", "Normal")
    location_style = header_plan.get("location_style", "Normal")
    contact_style = header_plan.get("contact_style", "Normal")
    order = header_plan.get(
        "line_order",
        ["name", "title", "department_or_practice", "location", "contact"],
    )

    # Build contact line (email | phone | linkedin)
    contact_bits = []
    if bio.email:
        contact_bits.append(bio.email)
    if bio.phone:
        contact_bits.append(bio.phone)
    if bio.linkedin_url:
        contact_bits.append(str(bio.linkedin_url))
    contact_line = " | ".join(contact_bits) if contact_bits else None

    for item in order:
        if item == "name" and bio.full_name:
            p = doc.add_paragraph(bio.full_name)
            _safe_set_style(p, name_style)
        elif item == "title" and bio.current_title:
            p = doc.add_paragraph(bio.current_title)
            _safe_set_style(p, title_style)
        elif item == "department_or_practice" and bio.department_or_practice:
            p = doc.add_paragraph(bio.department_or_practice)
            _safe_set_style(p, dept_style)
        elif item == "location" and bio.location:
            p = doc.add_paragraph(bio.location)
            _safe_set_style(p, location_style)
        elif item == "contact" and contact_line:
            p = doc.add_paragraph(contact_line)
            _safe_set_style(p, contact_style)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT


# ---------- main rendering ----------
def render_bio_to_docx(bio: Bio, plan: Dict[str, Any], out_path: Path):
    """
    Renders a .docx using the LLM-produced render plan:
      plan["header"]  -> header style mapping + order
      plan["sections"] -> list of section dicts with keys:
           key, label, heading_style, body_style, bullet_style, bullet_limit
      plan.get("tone") is informational (not enforced here)
    """
    doc = Document()

    # Header (LLM-driven)
    header_plan = plan.get("header", {}) or {}
    _render_header(doc, bio, header_plan)

    # Optional fallbacks (in case a section omits style names)
    legacy_styles = plan.get("styles", {}) or {}
    heading_fallback = legacy_styles.get("heading_style", "Heading 2")
    body_fallback = legacy_styles.get("body_style", "Normal")
    bullet_fallback = legacy_styles.get("bullet_style", "List Bullet")

    # Sections
    for section in plan.get("sections", []):
        label = section.get("label")
        key = section.get("key")
        if not label:
            # Skip unknown blocks without labels
            continue

        heading_style = section.get("heading_style", heading_fallback)
        body_style = section.get("body_style", body_fallback)
        bullet_style = section.get("bullet_style", bullet_fallback)
        bullet_limit = section.get("bullet_limit", 0)

        # Heading
        _add_heading(doc, label, heading_style)

        # Content
        if key == "summary":
            if bio.summary_paragraph:
                _add_paragraph(doc, bio.summary_paragraph, body_style)

        elif key == "expertise":
            _add_bullets(doc, bio.expertise_bullets, bullet_style, limit=bullet_limit)

        elif key == "experience":
            # Render each experience as a body line + optional impact bullet
            for item in bio.selected_experience:
                line = item.client_or_project
                if item.role:
                    line += f" — {item.role}"
                _add_paragraph(doc, line, body_style)
                if item.impact:
                    _add_bullets(doc, [item.impact], bullet_style, limit=1)

        elif key == "education":
            _add_bullets(doc, bio.education, bullet_style, limit=bullet_limit)

        elif key == "certifications":
            _add_bullets(doc, bio.certifications, bullet_style, limit=bullet_limit)

        elif key == "affiliations":
            _add_bullets(doc, bio.affiliations, bullet_style, limit=bullet_limit)

        elif key == "awards":
            _add_bullets(doc, bio.awards, bullet_style, limit=bullet_limit)

        else:
            # Custom/unknown section: if the Bio has a list attr with this name, render bullets
            vals = getattr(bio, key, None)
            if isinstance(vals, list):
                _add_bullets(doc, vals, bullet_style, limit=bullet_limit)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

import json
import re
from pathlib import Path
from typing import Any, Dict

def safe_load_json(path: Path) -> Dict[str, Any]:
    """Load JSON safely, even if file has BOM, extra logs, or mixed encodings."""
    raw = path.read_bytes()

    # strip BOM if present
    if raw.startswith(b"\xef\xbb\xbf"):
        raw = raw[3:]

    # try decoding with utf-8; fall back to latin-1
    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError:
        text = raw.decode("latin-1")

    text = text.strip()

    # try normal json first
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # try to find the first JSON object or array in the text (ignoring logs)
    m = re.search(r"(\{.*\}|\[.*\])", text, flags=re.S)
    if m:
        snippet = m.group(1)
        try:
            return json.loads(snippet)
        except json.JSONDecodeError:
            # remove stray commas before closing braces/brackets
            cleaned = re.sub(r",\s*([}\]])", r"\1", snippet)
            return json.loads(cleaned)

    # if nothing worked, show a preview
    preview = text[:300].replace("\n", "\\n")
    raise ValueError(f"File is not valid JSON. Preview:\n{preview}")



# ---------- CLI (optional) ----------
if __name__ == "__main__":
    import argparse

    ap = argparse.ArgumentParser(description="Render Bio to .docx using an LLM render plan")
    ap.add_argument("--plan", required=True, help="Path to render_plan.json produced by style_planner_ai")
    ap.add_argument("--bio-json", required=True, help="Path to Bio JSON (output of extractor_ai)")
    ap.add_argument("--out", default="output/standardized_bio.docx", help="Output .docx path")
    args = ap.parse_args()

    plan = load_plan(Path(args.plan))
    bio_data = safe_load_json(Path(args.bio_json))
    bio = Bio(**bio_data)

    render_bio_to_docx(bio, plan, Path(args.out))
    print(f"✓ Wrote {args.out}")
